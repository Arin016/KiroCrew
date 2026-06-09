"""Unit tests for the Knowledge Library (store, chunker, readers, extractor, retrieval)."""

from __future__ import annotations

import importlib
import json
import sys

import pytest

from kiro_claw.knowledge.chunker import HeadingAwareChunker
from kiro_claw.knowledge.extractor import EntityExtractor
from kiro_claw.knowledge.readers import FileReader
from kiro_claw.knowledge.retrieval import HybridRetriever, _bytes_to_floats
from kiro_claw.knowledge.store import KnowledgeStore, SimpleDiGraph

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def store(tmp_path):
    s = KnowledgeStore(str(tmp_path / "test.db"))
    yield s
    s.close()


@pytest.fixture()
def store_factory(tmp_path):
    """Return a callable that creates a new store at a given path."""
    stores = []

    def _make(name="test.db"):
        s = KnowledgeStore(str(tmp_path / name))
        stores.append(s)
        return s

    yield _make
    for s in stores:
        s.close()


# ---------------------------------------------------------------------------
# 1. KnowledgeStore
# ---------------------------------------------------------------------------

class TestKnowledgeStore:
    def test_create_and_get_item(self, store):
        item_id = store.add_item("Auth Design", "JWT tokens with 1h expiry", "design_doc",
                                 summary="Auth overview", tags=["auth", "jwt"])
        item = store.get_item(item_id)
        assert item is not None
        assert item["title"] == "Auth Design"
        assert item["content"] == "JWT tokens with 1h expiry"
        assert item["item_type"] == "design_doc"
        assert item["summary"] == "Auth overview"
        assert json.loads(item["tags"]) == ["auth", "jwt"]

    def test_fts_search(self, store):
        store.add_item("Auth Design", "JWT tokens with refresh flow", "design_doc")
        store.add_item("Database Schema", "DynamoDB table layout", "design_doc")
        results = store.search_items_fts("JWT")
        assert len(results) >= 1
        assert results[0]["title"] == "Auth Design"

    def test_add_entity_and_relation(self, store):
        e1 = store.add_entity("AuthService", "service", description="Handles auth")
        e2 = store.add_entity("DynamoDB", "technology", description="NoSQL DB")
        rid = store.add_entity_relation(e1, e2, "uses", description="Stores tokens")
        assert rid is not None
        assert store.graph.has_edge(e1, e2)
        edge = store.graph.edges[e1, e2]
        assert edge["relation_type"] == "uses"

    def test_entity_subgraph(self, store):
        e1 = store.add_entity("ServiceA", "service")
        e2 = store.add_entity("ServiceB", "service")
        e3 = store.add_entity("Database", "technology")
        store.add_entity_relation(e1, e2, "calls")
        store.add_entity_relation(e2, e3, "uses")
        sg = store.get_entity_subgraph(e1, depth=2)
        node_ids = {n["id"] for n in sg["nodes"]}
        assert e1 in node_ids
        assert e2 in node_ids
        assert e3 in node_ids
        assert len(sg["edges"]) == 2
        # Verify D3.js format: nodes have id/name/type, edges have source/target/type
        for n in sg["nodes"]:
            assert "id" in n and "name" in n and "type" in n
        for e in sg["edges"]:
            assert "source" in e and "target" in e and "type" in e

    def test_export_import_roundtrip(self, store_factory):
        s1 = store_factory("export.db")
        s1.add_item("Doc A", "Content A", "design_doc")
        s1.add_item("Doc B", "Content B", "runbook")
        s1.add_entity("SvcX", "service")
        bundle = s1.export_all()
        assert len(bundle["items"]) == 2
        assert len(bundle["entities"]) == 1

        s2 = store_factory("import.db")
        result = s2.import_bundle(bundle)
        assert result["items_imported"] == 2
        assert result["entities_created"] == 1
        stats = s2.get_stats()
        assert stats["items"] == 2
        assert stats["entities"] == 1

    def test_delete_item(self, store):
        item_id = store.add_item("Temp Doc", "Will be deleted", "personal_notes")
        assert store.get_item(item_id) is not None
        store.delete_item(item_id)
        assert store.get_item(item_id) is None
        # FTS should also be clean
        assert store.search_items_fts("deleted") == []

    def test_find_entity_case_insensitive(self, store):
        store.add_entity("DynamoDB", "technology")
        found = store.find_entity("dynamodb")
        assert found is not None
        assert found["name"] == "DynamoDB"

    def test_merge_entities(self, store):
        e_keep = store.add_entity("AuthService", "service")
        e_merge = store.add_entity("Auth Service", "service")
        e_other = store.add_entity("Database", "technology")
        store.add_entity_relation(e_merge, e_other, "uses")
        item_id = store.add_item("Doc", "content", "design_doc")
        store.add_mention(item_id, e_merge)

        store.merge_entities(e_keep, e_merge)

        # Merged entity should be gone
        assert store.find_entity("Auth Service") is None
        # Relation should point to kept entity
        rels = store.db.execute(
            "SELECT * FROM entity_relations WHERE source_id = ?", (e_keep,)
        ).fetchall()
        assert len(rels) == 1
        assert rels[0]["target_id"] == e_other
        # Mention should reference kept entity
        mentions = store.db.execute(
            "SELECT * FROM mentions WHERE entity_id = ?", (e_keep,)
        ).fetchall()
        assert len(mentions) == 1


# ---------------------------------------------------------------------------
# 2. HeadingAwareChunker
# ---------------------------------------------------------------------------

class TestHeadingAwareChunker:
    def test_chunk_markdown(self):
        text = "# Introduction\nThis is the intro paragraph.\n\n# Details\nHere are the details."
        chunker = HeadingAwareChunker(target_size=10)  # Very small to force split
        chunks = chunker.chunk(text)
        assert len(chunks) >= 2
        for c in chunks:
            assert "line_start" in c and "line_end" in c
            assert "content" in c
            assert c["chunk_index"] >= 0

    def test_chunk_code(self):
        code = "import os\n\ndef foo():\n    return 1\n\ndef bar():\n    return 2\n"
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk_code(code, language="python")
        assert len(chunks) >= 1
        # All code should be present across chunks
        combined = "\n".join(c["content"] for c in chunks)
        assert "def foo():" in combined
        assert "def bar():" in combined
        for c in chunks:
            assert "line_start" in c and "line_end" in c

    def test_small_text_single_chunk(self):
        text = "Just a short note."
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk(text)
        assert len(chunks) == 1
        assert chunks[0]["content"] == text


# ---------------------------------------------------------------------------
# 3. FileReader
# ---------------------------------------------------------------------------

class TestFileReader:
    def test_read_markdown(self, tmp_path):
        md = tmp_path / "test.md"
        md.write_text("# Hello\nWorld", encoding="utf-8")
        reader = FileReader()
        text, meta = reader.read(str(md))
        assert "# Hello" in text
        assert "World" in text
        assert meta["format"] == "md"
        assert meta["title"] == "test"
        assert meta["line_count"] == 2

    def test_read_unsupported(self, tmp_path):
        f = tmp_path / "data.xyz"
        f.write_text("binary-ish", encoding="utf-8")
        reader = FileReader()
        # Unsupported extension still falls through to _read_text
        text, meta = reader.read(str(f))
        assert "binary-ish" in text

    def test_supported_formats(self):
        reader = FileReader()
        for ext in ('.md', '.txt', '.py', '.html', '.json', '.yaml', '.csv'):
            assert ext in reader.SUPPORTED, f"{ext} missing from SUPPORTED"


# ---------------------------------------------------------------------------
# 4. EntityExtractor
# ---------------------------------------------------------------------------

class TestEntityExtractor:
    def test_extract_no_agent(self):
        import asyncio
        ext = EntityExtractor(pool=None)
        result = asyncio.get_event_loop().run_until_complete(ext.extract("some text"))
        assert result == {"title": "", "entities": [], "relations": [], "category": "document", "summary": ""}

    def test_parse_json_response(self):
        ext = EntityExtractor()
        raw = json.dumps({
            "entities": [{"name": "Svc", "type": "service", "description": "A service"}],
            "relations": [],
            "category": "design_doc",
            "summary": "A service doc."
        })
        result = ext._parse_response(raw)
        assert len(result["entities"]) == 1
        assert result["category"] == "design_doc"

    def test_parse_code_block_response(self):
        ext = EntityExtractor()
        raw = '```json\n{"entities": [], "relations": [], "category": "runbook", "summary": "ops"}\n```'
        result = ext._parse_response(raw)
        assert result["category"] == "runbook"
        assert result["summary"] == "ops"


# ---------------------------------------------------------------------------
# 5. HybridRetriever
# ---------------------------------------------------------------------------

class TestHybridRetriever:
    def test_keyword_search(self, store):
        store.add_item("Auth Design", "JWT tokens with refresh flow", "design_doc")
        store.add_item("DB Schema", "DynamoDB table layout", "design_doc")
        retriever = HybridRetriever(store)
        results = retriever.search("JWT")
        assert len(results) >= 1
        assert results[0]["title"] == "Auth Design"
        assert "keyword" in results[0]["match_type"]

    def test_rrf_fuse(self):
        list_a = [("item1", 1), ("item2", 2), ("item3", 3)]
        list_b = [("item2", 1), ("item3", 2), ("item4", 3)]
        fused = HybridRetriever._rrf_fuse(list_a, list_b, None, k=60)
        ids = [item_id for item_id, _ in fused]
        # item2 appears in both lists at good ranks, should be top
        assert ids[0] == "item2"
        # All 4 items should be present
        assert set(ids) == {"item1", "item2", "item3", "item4"}


# ---------------------------------------------------------------------------
# 6. SimpleDiGraph
# ---------------------------------------------------------------------------


class TestSimpleDiGraph:
    def test_add_node_and_has_node(self):
        g = SimpleDiGraph()
        g.add_node("a", name="A")
        assert g.has_node("a")
        assert not g.has_node("b")

    def test_add_edge_and_has_edge(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_edge("a", "b", weight=1.0)
        assert g.has_edge("a", "b")
        assert not g.has_edge("b", "a")

    def test_successors_predecessors(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_node("c")
        g.add_edge("a", "b")
        g.add_edge("a", "c")
        assert set(g.successors("a")) == {"b", "c"}
        assert set(g.predecessors("b")) == {"a"}
        assert list(g.successors("c")) == []

    def test_degree(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_node("b")
        g.add_node("c")
        g.add_edge("a", "b")
        g.add_edge("c", "a")
        assert g.degree("a") == 2  # 1 outgoing + 1 incoming
        assert g.degree("b") == 1

    def test_nodes_iteration_and_subscript(self):
        g = SimpleDiGraph()
        g.add_node("x", name="X", entity_type="svc")
        g.add_node("y", name="Y", entity_type="db")
        assert set(g.nodes) == {"x", "y"}
        assert g.nodes["x"]["name"] == "X"
        assert "x" in g.nodes
        assert len(g.nodes) == 2

    def test_edges_iteration_and_subscript(self):
        g = SimpleDiGraph()
        g.add_edge("a", "b", relation_type="calls")
        edges = list(g.edges(data=True))
        assert len(edges) == 1
        assert edges[0] == ("a", "b", {"relation_type": "calls"})
        assert g.edges["a", "b"]["relation_type"] == "calls"

    def test_clear(self):
        g = SimpleDiGraph()
        g.add_node("a")
        g.add_edge("a", "b")
        g.clear()
        assert not g.has_node("a")
        assert not g.has_edge("a", "b")
        assert list(g.nodes) == []


# ---------------------------------------------------------------------------
# 7. KnowledgeStore -- additional coverage
# ---------------------------------------------------------------------------


class TestKnowledgeStoreExtended:
    def test_update_item_fts_sync(self, store):
        item_id = store.add_item("Original", "old content about cats", "doc")
        assert len(store.search_items_fts("cats")) == 1
        store.update_item(item_id, title="Updated", content="new content about dogs")
        # After update, new content should be searchable
        assert len(store.search_items_fts("dogs")) == 1
        item = store.get_item(item_id)
        assert item["title"] == "Updated"
        assert item["content"] == "new content about dogs"

    def test_update_item_no_fields(self, store):
        item_id = store.add_item("Doc", "content", "doc")
        store.update_item(item_id)  # no-op, should not crash

    def test_update_item_non_fts_field(self, store):
        item_id = store.add_item("Doc", "content", "doc")
        store.update_item(item_id, status="archived")
        assert store.get_item(item_id)["status"] == "archived"

    def test_get_item_missing(self, store):
        assert store.get_item("nonexistent") is None

    def test_add_source_and_get_by_uri(self, store):
        sid = store.add_source("myfile", "local_file", "/tmp/test.md",
                               properties={"content_hash": "abc123"})
        found = store.get_source_by_uri("/tmp/test.md")
        assert found is not None
        assert found["id"] == sid
        assert store.get_source_by_uri("/tmp/nope") is None

    def test_update_source(self, store):
        sid = store.add_source("f", "local_file", "/tmp/f.md")
        store.update_source(sid, last_synced="2026-01-01T00:00:00")
        row = store.db.execute("SELECT last_synced FROM sources WHERE id = ?", (sid,)).fetchone()
        assert row["last_synced"] == "2026-01-01T00:00:00"

    def test_update_source_no_fields(self, store):
        sid = store.add_source("f", "local_file", "/tmp/f2.md")
        store.update_source(sid)  # no-op

    def test_add_source_location(self, store):
        sid = store.add_source("f", "local_file", "/tmp/loc.md")
        item_id = store.add_item("Doc", "content", "doc", source_id=sid)
        store.add_source_location(item_id, sid, chunk_range="0-10", section_title="Intro")
        rows = store.db.execute(
            "SELECT * FROM source_locations WHERE item_id = ?", (item_id,)).fetchall()
        assert len(rows) == 1
        assert rows[0]["section_title"] == "Intro"

    def test_get_neighbors_depth(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        e3 = store.add_entity("C", "svc")
        store.add_entity_relation(e1, e2, "calls")
        store.add_entity_relation(e2, e3, "calls")
        # depth=1 should get B only
        n1 = store.get_neighbors(e1, depth=1)
        assert {n["id"] for n in n1} == {e2}
        # depth=2 should get B and C
        n2 = store.get_neighbors(e1, depth=2)
        assert {n["id"] for n in n2} == {e2, e3}

    def test_get_neighbors_bidirectional(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        store.add_entity_relation(e2, e1, "calls")
        # e1 has no outgoing but has incoming from e2
        neighbors = store.get_neighbors(e1, depth=1)
        assert {n["id"] for n in neighbors} == {e2}

    def test_find_entity_by_alias(self, store):
        store.add_entity("DynamoDB", "technology", aliases=["ddb", "dynamo"])
        found = store.find_entity("ddb")
        assert found is not None
        assert found["name"] == "DynamoDB"

    def test_find_entity_not_found(self, store):
        assert store.find_entity("nonexistent") is None

    def test_export_item_with_entities(self, store):
        sid = store.add_source("f", "local_file", "/tmp/exp.md")
        item_id = store.add_item("Doc", "content", "doc", source_id=sid)
        e1 = store.add_entity("Svc", "service")
        e2 = store.add_entity("DB", "technology")
        store.add_mention(item_id, e1)
        store.add_mention(item_id, e2)
        store.add_entity_relation(e1, e2, "uses", source_item_id=item_id)
        store.add_source_location(item_id, sid, section_title="Main")
        bundle = store.export_item(item_id)
        assert bundle["item"]["id"] == item_id
        assert len(bundle["entities"]) == 2
        assert len(bundle["relations"]) == 1
        assert len(bundle["source_locations"]) == 1

    def test_export_item_missing(self, store):
        assert store.export_item("nope") == {}

    def test_delete_item_cleans_mentions(self, store):
        item_id = store.add_item("Doc", "content", "doc")
        eid = store.add_entity("Svc", "service")
        store.add_mention(item_id, eid, context="test")
        sid = store.add_source("f", "local_file", "/tmp/del.md")
        store.add_source_location(item_id, sid)
        store.delete_item(item_id)
        assert store.db.execute("SELECT * FROM mentions WHERE item_id = ?", (item_id,)).fetchone() is None
        assert store.db.execute("SELECT * FROM source_locations WHERE item_id = ?", (item_id,)).fetchone() is None

    def test_get_stats(self, store):
        store.add_item("A", "a", "doc")
        store.add_entity("E", "svc")
        stats = store.get_stats()
        assert stats["items"] == 1
        assert stats["entities"] == 1
        assert stats["relations"] == 0
        assert stats["sources"] == 0

    def test_graph_has_node(self, store):
        eid = store.add_entity("Svc", "service")
        assert store.graph.has_node(eid)
        assert not store.graph.has_node("fake")

    def test_graph_degree(self, store):
        e1 = store.add_entity("A", "svc")
        e2 = store.add_entity("B", "svc")
        store.add_entity_relation(e1, e2, "calls")
        assert store.graph.degree(e1) == 1
        assert store.graph.degree(e2) == 1

    def test_load_graph_on_reopen(self, tmp_path):
        db_path = str(tmp_path / "reload.db")
        s1 = KnowledgeStore(db_path)
        e1 = s1.add_entity("A", "svc")
        e2 = s1.add_entity("B", "svc")
        s1.add_entity_relation(e1, e2, "calls")
        s1.close()
        s2 = KnowledgeStore(db_path)
        assert s2.graph.has_node(e1)
        assert s2.graph.has_edge(e1, e2)
        s2.close()


# ---------------------------------------------------------------------------
# 8. HybridRetriever -- additional coverage
# ---------------------------------------------------------------------------


class TestHybridRetrieverExtended:
    def test_graph_search(self, store):
        e1 = store.add_entity("JWT", "concept")
        item_id = store.add_item("Auth", "JWT token design", "doc")
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("JWT")
        assert len(results) >= 1
        assert results[0][0] == item_id

    def test_graph_search_no_match(self, store):
        retriever = HybridRetriever(store)
        assert retriever._graph_search("nonexistent") == []

    def test_graph_search_with_neighbors(self, store):
        e1 = store.add_entity("Auth", "service")
        e2 = store.add_entity("JWT", "concept")
        store.add_entity_relation(e1, e2, "uses")
        item_id = store.add_item("Token doc", "about tokens", "doc")
        store.add_mention(item_id, e2)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("Auth")
        assert len(results) >= 1

    def test_vector_search_no_embedder(self, store):
        retriever = HybridRetriever(store, embedder=None)
        assert retriever._vector_search("query") is None

    def test_vector_search_with_embedder(self, store):
        emb = json.dumps([1.0, 0.0, 0.0])
        store.add_item("Vec Doc", "vector content", "doc", embedding=emb)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0, 0.0])
        results = retriever._vector_search("query")
        assert results is not None
        assert len(results) == 1

    def test_cosine_similarity_identical(self):
        assert HybridRetriever._cosine_similarity([1, 0], [1, 0]) == pytest.approx(1.0)

    def test_cosine_similarity_orthogonal(self):
        assert HybridRetriever._cosine_similarity([1, 0], [0, 1]) == pytest.approx(0.0)

    def test_cosine_similarity_zero_vector(self):
        assert HybridRetriever._cosine_similarity([0, 0], [1, 1]) == 0.0

    def test_search_combined_match_types(self, store):
        e1 = store.add_entity("JWT", "concept")
        emb = json.dumps([1.0, 0.0])
        item_id = store.add_item("JWT Auth", "JWT token design", "doc", embedding=emb)
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store, embedder=lambda q: [1.0, 0.0])
        results = retriever.search("JWT")
        assert len(results) >= 1
        # Should have multiple match types
        mt = results[0]["match_type"]
        assert "keyword" in mt

    def test_search_graph_pair_terms(self, store):
        """Graph search tries consecutive word pairs."""
        e1 = store.add_entity("Auth Service", "service")
        item_id = store.add_item("Doc", "about auth service", "doc")
        store.add_mention(item_id, e1)
        retriever = HybridRetriever(store)
        results = retriever._graph_search("Auth Service details")
        assert len(results) >= 1

    def test_bytes_to_floats_valid(self):
        assert _bytes_to_floats(json.dumps([1.0, 2.0]).encode()) == [1.0, 2.0]

    def test_bytes_to_floats_empty(self):
        assert _bytes_to_floats(b"") == []
        assert _bytes_to_floats(None) == []

    def test_bytes_to_floats_invalid(self):
        assert _bytes_to_floats(b"not json") == []


# ---------------------------------------------------------------------------
# 9. EntityExtractor -- additional coverage
# ---------------------------------------------------------------------------


class TestEntityExtractorExtended:
    def test_extract_empty_text(self):
        import asyncio
        ext = EntityExtractor(pool=None)
        result = asyncio.get_event_loop().run_until_complete(ext.extract(""))
        assert result == {"title": "", "entities": [], "relations": [], "category": "document", "summary": ""}

    def test_extract_with_agent(self):
        import asyncio

        class MockPool:
            async def send(self, prompt, timeout=60.0):
                return json.dumps({
                    "entities": [{"name": "Svc", "type": "service", "description": "A"}],
                    "relations": [], "category": "design_doc", "summary": "test"
                })

            async def send_batch(self, prompts, timeout=60.0):
                return [await self.send(p, timeout) for p in prompts]

        ext = EntityExtractor(pool=MockPool())
        result = asyncio.get_event_loop().run_until_complete(ext.extract("some text"))
        assert result["category"] == "design_doc"
        assert len(result["entities"]) == 1

    def test_extract_agent_exception(self):
        import asyncio

        class BadPool:
            async def send(self, prompt, timeout=60.0):
                raise RuntimeError("fail")

            async def send_batch(self, prompts, timeout=60.0):
                raise RuntimeError("fail")

        ext = EntityExtractor(pool=BadPool())
        result = asyncio.get_event_loop().run_until_complete(ext.extract("text"))
        assert result == {"title": "", "entities": [], "relations": [], "category": "document", "summary": ""}

    def test_parse_response_regex_fallback(self):
        ext = EntityExtractor()
        raw = 'Some preamble text {"entities": [], "relations": [], "category": "runbook", "summary": "ok"} trailing'
        result = ext._parse_response(raw)
        assert result["category"] == "runbook"

    def test_parse_response_garbage(self):
        ext = EntityExtractor()
        result = ext._parse_response("totally invalid garbage")
        assert result == {"title": "", "entities": [], "relations": [], "category": "document", "summary": ""}

    def test_extract_code_block(self):
        ext = EntityExtractor()
        assert ext._extract_code_block("no block here") is None
        result = ext._extract_code_block('```\n{"a": 1}\n```')
        assert result == '{"a": 1}'

    def test_validate_partial_data(self):
        ext = EntityExtractor()
        result = ext._validate({"category": "runbook"})
        assert result["entities"] == []
        assert result["relations"] == []
        assert result["summary"] == ""
        assert result["category"] == "runbook"


# ---------------------------------------------------------------------------
# 10. Chunker -- additional coverage
# ---------------------------------------------------------------------------


class TestChunkerExtended:
    def test_chunk_with_overlap(self):
        text = "# A\n" + " ".join(["word"] * 600) + "\n\n# B\n" + " ".join(["other"] * 100)
        chunker = HeadingAwareChunker(target_size=500, overlap=10)
        chunks = chunker.chunk(text)
        assert len(chunks) >= 2
        # Second chunk should contain overlap from first
        if len(chunks) > 1:
            assert chunks[1]["chunk_index"] == 1

    def test_chunk_slides(self):
        text = "## Slide 1: Intro\nHello world\n\n## Slide 2: Details\nMore info"
        chunker = HeadingAwareChunker()
        slides = chunker.chunk_slides(text)
        assert len(slides) == 2
        assert slides[0]["section_title"] == "Slide 1: Intro"
        assert "Hello world" in slides[0]["content"]

    def test_chunk_code_oversized(self):
        # Generate a single huge function
        lines = ["def big():"] + [f"    x = {i}" for i in range(1000)]
        code = "\n".join(lines)
        chunker = HeadingAwareChunker(target_size=50)
        chunks = chunker.chunk_code(code, language="python")
        assert len(chunks) > 1
        combined = "\n".join(c["content"] for c in chunks)
        assert "def big():" in combined

    def test_chunk_no_headings(self):
        text = "Just plain text without any headings at all."
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk(text)
        assert len(chunks) == 1
        assert chunks[0]["section_title"] is None

    def test_small_overlap_does_not_duplicate_whole_prev_chunk(self):
        """Regression: overlap=1 passed the ``overlap > 0`` guard but
        ``int(1 / 1.3) == 0``, and ``prev_words[-0:]`` is ``prev_words[0:]`` — the
        ENTIRE previous chunk. So a small (but valid, user-configurable via a source's
        ``chunk_overlap`` property) overlap silently prepended the whole previous chunk
        to every subsequent chunk, duplicating content across the knowledge base.
        """
        words = " ".join(f"w{i}" for i in range(400))
        chunker = HeadingAwareChunker(target_size=50, overlap=1)
        chunks = chunker.chunk(words)
        assert len(chunks) >= 2  # must actually split to exercise the overlap path

        prev_word_count = len(chunks[0]["content"].split())
        # The overlap prefix is the first line of chunk[1] (joined with "\n" + content).
        overlap_prefix = chunks[1]["content"].split("\n", 1)[0]
        injected = len(overlap_prefix.split())
        # A tiny overlap must inject a tiny prefix — never (almost) the whole prev chunk.
        assert injected < prev_word_count, (
            f"overlap=1 injected {injected} words but previous chunk has "
            f"{prev_word_count} — the entire previous chunk was duplicated"
        )
        assert injected <= 2, f"overlap=1 should inject ~0-1 words, got {injected}"

    def test_zero_overlap_injects_nothing(self):
        words = " ".join(f"w{i}" for i in range(400))
        chunker = HeadingAwareChunker(target_size=50, overlap=0)
        chunks = chunker.chunk(words)
        assert len(chunks) >= 2
        # With overlap=0 chunk[1] has no injected prefix line from chunk[0].
        assert not chunks[1]["content"].startswith(chunks[0]["content"].split()[0] + " w")

    def test_large_overlap_still_works(self):
        # The fix must not change behavior for the normal/default overlap.
        words = " ".join(f"w{i}" for i in range(400))
        chunker = HeadingAwareChunker(target_size=50, overlap=200)
        chunks = chunker.chunk(words)
        assert len(chunks) >= 2
        overlap_prefix = chunks[1]["content"].split("\n", 1)[0]
        # int(200/1.3) = 153, capped by prev chunk length — a real, multi-word overlap.
        assert len(overlap_prefix.split()) >= 2


# ---------------------------------------------------------------------------
# 11. FileReader -- additional coverage
# ---------------------------------------------------------------------------


class TestFileReaderExtended:
    def test_read_html_without_html2text(self, tmp_path):
        """Test HTML reading (exercises html2text or regex fallback)."""
        html_file = tmp_path / "test.html"
        html_file.write_text("<html><body><p>Hello</p></body></html>")
        reader = FileReader()
        text, meta = reader.read(str(html_file))
        assert "Hello" in text

    def test_read_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin.txt"
        f.write_bytes(b"caf\xe9")
        reader = FileReader()
        text, meta = reader.read(str(f))
        assert "caf" in text

    def test_read_json_file(self, tmp_path):
        f = tmp_path / "data.json"
        f.write_text('{"key": "value"}')
        reader = FileReader()
        text, meta = reader.read(str(f))
        assert '"key"' in text
        assert meta["format"] == "json"


class TestPysqlite3Fallback:
    """Verify modules fall back to stdlib sqlite3 when pysqlite3 is unavailable."""

    _MODULES = (
        "kiro_claw.knowledge.store",
        "kiro_claw.knowledge.retrieval",
        "kiro_claw.snapshot",
    )

    def _reload_without_pysqlite3(self, module_name: str):
        """Force-reimport a module with pysqlite3 blocked."""
        import sqlite3 as stdlib_sqlite3

        saved = sys.modules.pop("pysqlite3", None)
        for mod in list(sys.modules):
            if mod == module_name or mod.startswith(module_name + "."):
                sys.modules.pop(mod)

        sys.modules["pysqlite3"] = None  # type: ignore[assignment]
        try:
            mod = importlib.import_module(module_name)
            assert mod.sqlite3 is stdlib_sqlite3
        finally:
            del sys.modules["pysqlite3"]
            if saved is not None:
                sys.modules["pysqlite3"] = saved

    def test_store_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("kiro_claw.knowledge.store")

    def test_retrieval_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("kiro_claw.knowledge.retrieval")

    def test_snapshot_falls_back_to_stdlib_sqlite3(self):
        self._reload_without_pysqlite3("kiro_claw.snapshot")


# ---------------------------------------------------------------------------
# 12. chunk_markdown() -- heading-aware markdown chunking
# ---------------------------------------------------------------------------


class TestChunkMarkdown:
    def test_splits_on_headings(self):
        text = "# Intro\nParagraph one.\n\n## Details\n" + " ".join(["detail"] * 400) + "\n\n## Conclusion\n" + " ".join(["final"] * 400)
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk_markdown(text)
        assert len(chunks) >= 2
        # All chunks have required fields
        for c in chunks:
            assert "content" in c
            assert "section_title" in c
            assert "chunk_index" in c
            assert "line_start" in c

    def test_preserves_markdown_formatting(self):
        text = "# Title\n\n**Bold text** and `code`.\n\n- List item 1\n- List item 2"
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk_markdown(text)
        combined = "\n".join(c["content"] for c in chunks)
        assert "**Bold text**" in combined
        assert "`code`" in combined
        assert "- List item" in combined

    def test_no_headings_falls_back(self):
        text = "Just plain text without headings."
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk_markdown(text)
        assert len(chunks) == 1
        assert chunks[0]["content"] == text

    def test_section_titles_extracted(self):
        text = "## Architecture\n" + " ".join(["arch"] * 300) + "\n\n## Security\n" + " ".join(["sec"] * 300)
        chunker = HeadingAwareChunker(target_size=500)
        chunks = chunker.chunk_markdown(text)
        titles = [c["section_title"] for c in chunks]
        assert "Architecture" in titles
        assert "Security" in titles

    def test_oversized_section_splits(self):
        text = "# Big Section\n" + " ".join(["word"] * 1000)
        chunker = HeadingAwareChunker(target_size=50)
        chunks = chunker.chunk_markdown(text)
        assert len(chunks) > 1


# ---------------------------------------------------------------------------
# 13. FileReader -- .docx content_type metadata
# ---------------------------------------------------------------------------


class TestDocxContentType:
    def test_docx_returns_content_type_markdown(self, tmp_path):
        """Verify _read_docx sets content_type: markdown in metadata."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        # Create a minimal .docx
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Some content here.")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        reader = FileReader()
        text, meta = reader.read(str(path))
        assert meta.get("content_type") == "markdown"
        assert "# Test Heading" in text
        assert "Some content here." in text

    def test_docx_content_type_in_dispatch(self):
        """Verify .docx is in the dispatch table."""
        reader = FileReader()
        assert '.docx' in reader._DISPATCH


class TestCosineSimilarityDimensionMismatch:
    """Regression: HybridRetriever._cosine_similarity must treat vectors of
    different dimensionality as incomparable (return 0.0), not silently truncate.

    The query vector is freshly embedded while the item vector is read from the DB,
    so a change in embedding dimensionality between ingestion and query yields
    mismatched lengths. With a plain ``zip(a, b)`` the dot product silently
    truncates to the shorter length while the norms still use the full vectors,
    producing a meaningless (often falsely high) similarity. The sibling code in
    ``vector_memory.py`` already guards this exact case (``if n_floats != q_len:
    continue``); this helper must not score across mismatched dimensions either.
    """

    def test_mismatched_dims_return_zero_not_false_match(self):
        # 8-dim query vs 4-dim stored item that happens to equal the query's prefix.
        # Truncating zip() makes these look identical (1.0); they are incomparable.
        query_vec = [1.0, 1.0, 1.0, 1.0, 0.0, 0.0, 0.0, 0.0]
        item_vec = [1.0, 1.0, 1.0, 1.0]
        sim = HybridRetriever._cosine_similarity(query_vec, item_vec)
        assert sim == 0.0, (
            f"mismatched-dimension vectors must be incomparable (0.0), got {sim} "
            "— dot product silently truncated while norms used full vectors"
        )

    def test_mismatched_dims_other_order_also_zero(self):
        # Order must not matter: shorter query vs longer item is equally incomparable.
        sim = HybridRetriever._cosine_similarity([1.0, 1.0, 1.0, 1.0],
                                                 [1.0, 1.0, 1.0, 1.0, 0.0, 0.0, 0.0, 0.0])
        assert sim == 0.0

    def test_equal_dims_unaffected(self):
        # The fix must not change behavior for the normal equal-length case.
        a = [1.0, 0.0, 0.0]
        b = [1.0, 0.0, 0.0]
        assert HybridRetriever._cosine_similarity(a, b) == pytest.approx(1.0)
        orthogonal = HybridRetriever._cosine_similarity([1.0, 0.0], [0.0, 1.0])
        assert orthogonal == pytest.approx(0.0)

    def test_empty_vector_edge_cases(self):
        # Document guard precedence (review nit): the length check runs first.
        # [] vs [1.0] are mismatched dims -> 0.0 (length guard wins).
        assert HybridRetriever._cosine_similarity([], [1.0]) == 0.0
        assert HybridRetriever._cosine_similarity([1.0], []) == 0.0
        # [] vs [] are equal-length but zero-norm -> 0.0 (zero-norm guard).
        assert HybridRetriever._cosine_similarity([], []) == 0.0
